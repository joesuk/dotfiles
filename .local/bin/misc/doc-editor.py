# this is a python script to find and replace attributes ($name and $amount) in a word document with entries taken from a csv and then separately generate different docx files for each row of the csv
from docx import Document
import csv

# Define lists to store names and amounts
titles = []
firstNames = []
secondNames = []
amounts = []

# Read the CSV file
with open('test.csv', newline='', encoding='utf-8') as csvfile:
    reader = csv.reader(csvfile)
    header = next(reader)  # Skip the header row if present
    for row in reader:
        if len(row) >= 2:  # Ensure there are at least two columns
            titles.append(row[0].strip())
            firstNames.append(row[1].strip())
            secondNames.append(row[2].strip())
            amounts.append(row[3].strip())

            print(row[0])
            print(row[1])

for i in range(len(names)):
    doc = Document("contribution.docx")
    for paragraph in doc.paragraphs:
        if '$name' in paragraph.text:
            fullName = titles[i] + ". " + firstNames[i] + " " + secondNames[i]
            paragraph.text = paragraph.text.replace("$name", fullName)
    for paragraph in doc.paragraphs:
        if '$amount' in paragraph.text:
            paragraph.text = paragraph.text.replace("$amount", amounts[i])
    doc.save(str(names[i]) + ".docx")
